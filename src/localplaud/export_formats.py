"""Portable recording exports for the Web App."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from xml.sax.saxutils import escape

from .config import get_settings
from .db.models import PlaudFile, StageName
from .db.session import session_scope
from .store.speakers import display_names

_PDF_FONT_PATH = Path(__file__).parent / "assets" / "fonts" / "NotoSansTC.ttf"


class MissingExportContentError(LookupError):
    """The selected recording has no current content for the requested export."""


def _plain_markdown_inline(value: str) -> str:
    value = re.sub(r"\[([^]]+)]\(([^)]+)\)", r"\1 (\2)", value)
    return re.sub(r"(\*\*|__|`|~~)", "", value).strip()


def _markdown_blocks(value: str) -> list[tuple[str, str]]:
    """Small, deterministic Markdown projection for portable document exports."""
    blocks: list[tuple[str, str]] = []
    in_code = False
    for raw_line in value.replace("\r\n", "\n").split("\n"):
        line = raw_line.rstrip()
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if not line.strip():
            if blocks and blocks[-1][0] != "space":
                blocks.append(("space", ""))
            continue
        if in_code:
            blocks.append(("code", line))
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            blocks.append((f"h{len(heading.group(1))}", _plain_markdown_inline(heading.group(2))))
            continue
        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if bullet:
            blocks.append(("bullet", _plain_markdown_inline(bullet.group(1))))
            continue
        numbered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if numbered:
            blocks.append(("number", _plain_markdown_inline(numbered.group(1))))
            continue
        blocks.append(("body", _plain_markdown_inline(line)))
    return blocks


def _stamp(seconds: float, *, milliseconds: bool = False) -> str:
    milliseconds_total = max(0, round(float(seconds or 0) * 1000))
    hours, remainder = divmod(milliseconds_total, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, millis = divmod(remainder, 1000)
    if milliseconds:
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
    if hours:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def recording_data(file_id: str) -> dict:
    with session_scope() as session:
        file = session.get(PlaudFile, file_id)
        if file is None:
            raise ValueError(f"no such file: {file_id}")
        settings = get_settings()
        if settings.pipeline.artifact_mode == "independent":
            raw = file.local_transcript
        elif settings.pipeline.prefer_cloud_artifacts:
            raw = file.plaud_transcript or file.local_transcript
        else:
            raw = file.local_transcript
        corrected = file.corrected_transcript_for_source(raw.source) if raw else None
        segments = list((corrected.segments if corrected else raw.segments) or []) if raw else []
        names = display_names(session, file.id)
        stale = {
            run.stage for run in file.stage_runs if (run.detail or {}).get("stale")
        }
        summaries = [
            summary
            for summary in file.summaries
            if (settings.pipeline.artifact_mode != "independent" or summary.source == "local")
            and summary.template != "mind_map"
            and not (
                summary.source == "local"
                and StageName.summarize in stale
            )
        ]
        notes = [
            {
                "title": (summary.template_snapshot or {}).get("name")
                or summary.title
                or summary.template.replace("-", " ").title(),
                "content": summary.content_md,
            }
            for summary in summaries
        ] + [
            {"title": note.title, "content": note.content_md} for note in file.user_notes
        ]
        return {
            "title": file.display_title,
            "segments": segments,
            "speaker_names": names,
            "notes": notes,
            "audio_path": file.audio_path,
            "transcript_provenance": (
                {
                    "transcript_id": raw.id,
                    "transcript_source": raw.source,
                    "transcript_revision_id": corrected.id if corrected else None,
                    "transcript_revision": corrected.revision if corrected else None,
                }
                if raw
                else {}
            ),
        }


def transcript_provenance(file_id: str) -> dict:
    """Return the canonical transcript lineage used by transcript exports."""
    return recording_data(file_id)["transcript_provenance"]


def render_transcript(
    file_id: str,
    fmt: str,
    *,
    timestamps: bool = True,
    speakers: bool = True,
) -> tuple[bytes, str]:
    return render_transcript_data(
        recording_data(file_id),
        fmt,
        timestamps=timestamps,
        speakers=speakers,
    )


def render_transcript_data(
    data: dict,
    fmt: str,
    *,
    timestamps: bool = True,
    speakers: bool = True,
) -> tuple[bytes, str]:
    """Render a transcript from one immutable recording-data snapshot."""
    segments = data["segments"]
    if not segments:
        raise MissingExportContentError("recording has no exportable transcript")

    def text_for(segment: dict) -> str:
        text = str(segment.get("text") or "").strip()
        speaker = segment.get("speaker")
        if speakers and speaker:
            text = f"{data['speaker_names'].get(speaker, speaker)}: {text}"
        return text

    def parts_for(segment: dict) -> tuple[str, str]:
        labels: list[str] = []
        if timestamps:
            labels.append(f"[{_stamp(segment.get('start') or 0)}]")
        speaker = segment.get("speaker")
        if speakers and speaker:
            labels.append(data["speaker_names"].get(speaker, speaker))
        return " · ".join(labels), str(segment.get("text") or "").strip()

    if fmt in {"srt", "vtt"}:
        blocks = []
        for index, segment in enumerate(segments, 1):
            start = _stamp(segment.get("start") or 0, milliseconds=True)
            end = _stamp(segment.get("end") or segment.get("start") or 0, milliseconds=True)
            if fmt == "vtt":
                start, end = start.replace(",", "."), end.replace(",", ".")
            blocks.append(f"{index}\n{start} --> {end}\n{text_for(segment)}")
        prefix = "WEBVTT\n\n" if fmt == "vtt" else ""
        return (prefix + "\n\n".join(blocks) + "\n").encode(), f"text/{fmt}"

    lines = []
    for segment in segments:
        prefix = f"[{_stamp(segment.get('start') or 0)}] " if timestamps else ""
        lines.append(prefix + text_for(segment))
    title = data["title"]
    if fmt == "txt":
        return (f"{title}\n\n" + "\n".join(lines) + "\n").encode(), "text/plain"
    if fmt == "docx":
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.page_width, section.page_height = Inches(8.5), Inches(11)
        section.top_margin = section.right_margin = Inches(1)
        section.bottom_margin = section.left_margin = Inches(1)
        section.header_distance = section.footer_distance = Inches(0.492)

        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(11)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25

        header = section.header.paragraphs[0]
        header.text = "localplaud · Transcript"
        header.runs[0].font.name = "Arial"
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = RGBColor(102, 112, 133)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run("Page ")
        footer_run.font.size = Pt(9)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)

        title_paragraph = document.add_paragraph()
        title_paragraph.paragraph_format.space_after = Pt(5)
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.name = "Arial"
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = RGBColor(11, 37, 69)
        subtitle = document.add_paragraph("Canonical transcript · exported by localplaud")
        subtitle.paragraph_format.space_after = Pt(18)
        subtitle.runs[0].font.size = Pt(10)
        subtitle.runs[0].font.color.rgb = RGBColor(102, 112, 133)

        for segment in segments:
            label, text = parts_for(segment)
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.keep_together = True
            if label:
                label_run = paragraph.add_run(label + "  ")
                label_run.bold = True
                label_run.font.color.rgb = RGBColor(31, 77, 120)
            paragraph.add_run(text)

        output = BytesIO()
        document.save(output)
        return output.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fmt == "pdf":
        from reportlab.lib.colors import HexColor
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.pdfbase.pdfmetrics import registerFont, registerFontFamily
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        registerFont(TTFont("NotoSansTC", str(_PDF_FONT_PATH)))
        registerFontFamily(
            "NotoSansTC",
            normal="NotoSansTC",
            bold="NotoSansTC",
            italic="NotoSansTC",
            boldItalic="NotoSansTC",
        )
        output = BytesIO()
        document = SimpleDocTemplate(
            output,
            pagesize=letter,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=0.82 * inch,
            bottomMargin=0.82 * inch,
            title=title,
            author="localplaud",
        )
        body = ParagraphStyle(
            "TranscriptBody",
            fontName="NotoSansTC",
            fontSize=11,
            leading=15,
            textColor=HexColor("#1F2937"),
            spaceAfter=8,
            alignment=TA_LEFT,
            splitLongWords=True,
        )
        title_style = ParagraphStyle(
            "TranscriptTitle",
            parent=body,
            fontSize=22,
            leading=27,
            textColor=HexColor("#0B2545"),
            spaceAfter=5,
        )
        subtitle_style = ParagraphStyle(
            "TranscriptSubtitle",
            parent=body,
            fontSize=9,
            leading=12,
            textColor=HexColor("#667085"),
            spaceAfter=18,
        )
        story = [
            Paragraph(escape(title), title_style),
            Paragraph("Canonical transcript · exported by localplaud", subtitle_style),
        ]
        for segment in segments:
            label, text = parts_for(segment)
            content = f'<font color="#1F4D78"><b>{escape(label)}</b></font>  ' if label else ""
            story.append(Paragraph(content + escape(text).replace("\n", "<br/>"), body))
            story.append(Spacer(1, 1))

        def decorate_page(canvas, doc):
            canvas.saveState()
            canvas.setFont("NotoSansTC", 8)
            canvas.setFillColor(HexColor("#667085"))
            canvas.drawString(inch, 10.45 * inch, "localplaud · Transcript")
            canvas.drawRightString(7.5 * inch, 0.48 * inch, f"Page {doc.page}")
            canvas.restoreState()

        document.build(story, onFirstPage=decorate_page, onLaterPages=decorate_page)
        return output.getvalue(), "application/pdf"
    raise ValueError("unsupported transcript format")


def render_notes(file_id: str, fmt: str) -> tuple[bytes, str]:
    return render_notes_data(recording_data(file_id), fmt)


def render_notes_data(data: dict, fmt: str) -> tuple[bytes, str]:
    """Render current notes from the same snapshot used for manifest lineage."""
    if not data["notes"]:
        raise MissingExportContentError("recording has no exportable notes")
    lines: list[str] = []
    markdown: list[str] = [f"# {data['title']}", ""]
    for note in data["notes"]:
        lines += [note["title"], note["content"], ""]
        markdown += [f"## {note['title']}", "", note["content"], ""]
    if fmt == "md":
        return "\n".join(markdown).encode(), "text/markdown"
    if fmt == "txt":
        return (data["title"] + "\n\n" + "\n".join(lines)).encode(), "text/plain"
    if fmt == "docx":
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.page_width, section.page_height = Inches(8.5), Inches(11)
        section.top_margin = section.right_margin = Inches(1)
        section.bottom_margin = section.left_margin = Inches(1)
        section.header_distance = section.footer_distance = Inches(0.492)
        for style_name, size, color, before, after in (
            ("Normal", 11, "1F2937", 0, 6),
            ("Heading 1", 16, "2E74B5", 18, 10),
            ("Heading 2", 13, "2E74B5", 14, 7),
            ("Heading 3", 12, "1F4D78", 10, 5),
        ):
            style = document.styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.line_spacing = 1.25
        for style_name in ("List Bullet", "List Number"):
            style = document.styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(11)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
            style.paragraph_format.left_indent = Inches(0.375)
            style.paragraph_format.first_line_indent = Inches(-0.188)
            style.paragraph_format.space_after = Pt(4)
            style.paragraph_format.line_spacing = 1.25

        header = section.header.paragraphs[0]
        header.text = "localplaud · Notes"
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = RGBColor(102, 112, 133)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run("Page ").font.size = Pt(9)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)

        title_paragraph = document.add_paragraph()
        title_paragraph.paragraph_format.space_after = Pt(5)
        title_run = title_paragraph.add_run(data["title"])
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = RGBColor(11, 37, 69)
        subtitle = document.add_paragraph("Notes · exported by localplaud")
        subtitle.paragraph_format.space_after = Pt(18)
        subtitle.runs[0].font.size = Pt(10)
        subtitle.runs[0].font.color.rgb = RGBColor(102, 112, 133)

        for note in data["notes"]:
            document.add_paragraph(note["title"], style="Heading 1")
            for kind, text in _markdown_blocks(note["content"]):
                if kind == "space":
                    continue
                if kind in {"h1", "h2", "h3"}:
                    document.add_paragraph(text, style=f"Heading {min(int(kind[1]) + 1, 3)}")
                elif kind == "bullet":
                    document.add_paragraph(text, style="List Bullet")
                elif kind == "number":
                    document.add_paragraph(text, style="List Number")
                else:
                    paragraph = document.add_paragraph(text)
                    if kind == "code":
                        paragraph.paragraph_format.left_indent = Inches(0.25)
                        paragraph.runs[0].font.name = "Menlo"
                        paragraph.runs[0].font.size = Pt(9)

        output = BytesIO()
        document.save(output)
        return output.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fmt == "pdf":
        from reportlab.lib.colors import HexColor
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.pdfbase.pdfmetrics import registerFont, registerFontFamily
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        registerFont(TTFont("NotoSansTC", str(_PDF_FONT_PATH)))
        registerFontFamily(
            "NotoSansTC",
            normal="NotoSansTC",
            bold="NotoSansTC",
            italic="NotoSansTC",
            boldItalic="NotoSansTC",
        )
        output = BytesIO()
        document = SimpleDocTemplate(
            output,
            pagesize=letter,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=0.82 * inch,
            bottomMargin=0.82 * inch,
            title=data["title"],
            author="localplaud",
        )
        body = ParagraphStyle(
            "NotesBody",
            fontName="NotoSansTC",
            fontSize=11,
            leading=15,
            textColor=HexColor("#1F2937"),
            spaceAfter=8,
        )
        styles = {
            "title": ParagraphStyle(
                "NotesTitle", parent=body, fontSize=22, leading=27,
                textColor=HexColor("#0B2545"), spaceAfter=5,
            ),
            "subtitle": ParagraphStyle(
                "NotesSubtitle", parent=body, fontSize=9, leading=12,
                textColor=HexColor("#667085"), spaceAfter=18,
            ),
            "h1": ParagraphStyle(
                "NotesH1", parent=body, fontSize=16, leading=20,
                textColor=HexColor("#2E74B5"), spaceBefore=18, spaceAfter=10,
            ),
            "h2": ParagraphStyle(
                "NotesH2", parent=body, fontSize=13, leading=17,
                textColor=HexColor("#2E74B5"), spaceBefore=14, spaceAfter=7,
            ),
            "h3": ParagraphStyle(
                "NotesH3", parent=body, fontSize=12, leading=16,
                textColor=HexColor("#1F4D78"), spaceBefore=10, spaceAfter=5,
            ),
            "code": ParagraphStyle(
                "NotesCode", parent=body, fontSize=9, leading=13,
                leftIndent=18, textColor=HexColor("#475467"),
                backColor=HexColor("#F2F4F7"), borderPadding=5,
            ),
            "list": ParagraphStyle(
                "NotesList", parent=body, leftIndent=27, firstLineIndent=-14,
                spaceAfter=4,
            ),
        }
        story = [
            Paragraph(escape(data["title"]), styles["title"]),
            Paragraph("Notes · exported by localplaud", styles["subtitle"]),
        ]
        for note in data["notes"]:
            story.append(Paragraph(escape(note["title"]), styles["h1"]))
            number = 0
            for kind, text in _markdown_blocks(note["content"]):
                if kind == "space":
                    story.append(Spacer(1, 3))
                elif kind in {"h1", "h2", "h3"}:
                    story.append(Paragraph(escape(text), styles[f"h{min(int(kind[1]) + 1, 3)}"]))
                elif kind == "bullet":
                    story.append(Paragraph(escape(text), styles["list"], bulletText="•"))
                elif kind == "number":
                    number += 1
                    story.append(Paragraph(escape(text), styles["list"], bulletText=f"{number}."))
                else:
                    story.append(Paragraph(escape(text), styles["code"] if kind == "code" else body))

        def decorate_page(canvas, doc):
            canvas.saveState()
            canvas.setFont("NotoSansTC", 8)
            canvas.setFillColor(HexColor("#667085"))
            canvas.drawString(inch, 10.45 * inch, "localplaud · Notes")
            canvas.drawRightString(7.5 * inch, 0.48 * inch, f"Page {doc.page}")
            canvas.restoreState()

        document.build(story, onFirstPage=decorate_page, onLaterPages=decorate_page)
        return output.getvalue(), "application/pdf"
    raise ValueError("unsupported notes format")
